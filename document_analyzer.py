"""
Chat&Talk GPT - Document Analyzer
PDF, Word, and File Analysis
"""
import os
import asyncio
import logging
from typing import Dict, Any, List, Optional
from datetime import datetime
import json
import re

logger = logging.getLogger("DocumentAnalyzer")

# Try importing document processing libraries
try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 not available")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


class DocumentAnalyzer:
    """
    Analyze documents (PDF, Word, Excel, Text)
    Extract text, summarize, and answer questions
    """
    
    def __init__(self):
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.supported_formats = {
            ".pdf": self._analyze_pdf,
            ".docx": self._analyze_docx,
            ".doc": self._analyze_docx,
            ".txt": self._analyze_text,
            ".csv": self._analyze_csv,
            ".xlsx": self._analyze_excel,
            ".xls": self._analyze_excel,
            ".pptx": self._analyze_pptx,
            ".md": self._analyze_text,
            ".json": self._analyze_text
        }
    
    async def analyze_file(
        self,
        file_path: str,
        file_content: bytes = None
    ) -> Dict[str, Any]:
        """
        Analyze a file and extract content
        
        Args:
            file_path: Path to file or filename
            file_content: Optional file bytes
            
        Returns:
            Dict with extracted content and metadata
        """
        logger.info(f"Analyzing file: {file_path}")
        
        # Get file extension
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.supported_formats:
            return {
                "success": False,
                "error": f"Unsupported file format: {ext}",
                "supported_formats": list(self.supported_formats.keys())
            }
        
        try:
            # Get content
            if file_content:
                content = file_content
            else:
                with open(file_path, "rb") as f:
                    content = f.read()
            
            # Check file size
            if len(content) > self.max_file_size:
                return {
                    "success": False,
                    "error": f"File too large. Max size: {self.max_file_size // (1024*1024)}MB"
                }
            
            # Analyze based on format
            analyzer = self.supported_formats[ext]
            result = await analyzer(content, file_path)
            
            result["filename"] = os.path.basename(file_path)
            result["file_size"] = len(content)
            result["file_type"] = ext
            result["analyzed_at"] = datetime.now().isoformat()
            
            return result
            
        except Exception as e:
            logger.error(f"File analysis error: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _analyze_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Analyze PDF file"""
        if not PYPDF2_AVAILABLE:
            return {
                "success": False,
                "error": "PDF processing not available. Install PyPDF2"
            }
        
        loop = asyncio.get_event_loop()
        
        def extract_pdf():
            from io import BytesIO
            pdf_file = BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text = ""
            pages = []
            metadata = {}
            
            # Extract metadata
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", "")
                }
            
            # Extract text from each page
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text += page_text + "\n\n"
                pages.append({
                    "page_number": i + 1,
                    "text": page_text[:1000]  # First 1000 chars
                })
            
            return {
                "text": text,
                "pages": pages,
                "page_count": len(reader.pages),
                "metadata": metadata
            }
        
        result = await loop.run_in_executor(None, extract_pdf)
        
        # Generate summary
        summary = self._generate_summary(result["text"])
        keywords = self._extract_keywords(result["text"])
        
        return {
            "success": True,
            "text": result["text"][:50000],  # Limit text
            "summary": summary,
            "keywords": keywords,
            "page_count": result["page_count"],
            "pages": result["pages"][:10],  # First 10 pages
            "metadata": result["metadata"]
        }
    
    async def _analyze_docx(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Analyze Word document"""
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "Word processing not available. Install python-docx"
            }
        
        loop = asyncio.get_event_loop()
        
        def extract_docx():
            from io import BytesIO
            doc_file = BytesIO(content)
            doc = Document(doc_file)
            
            text = ""
            paragraphs = []
            tables = []
            
            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text += para.text + "\n"
                    paragraphs.append({
                        "number": i + 1,
                        "text": para.text[:500]
                    })
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "table_number": i + 1,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                    "data": table_data[:10]  # First 10 rows
                })
                text += f"\n[Table {i+1}]\n" + "\n".join([", ".join(row) for row in table_data[:5]])
            
            return {
                "text": text,
                "paragraph_count": len(paragraphs),
                "paragraphs": paragraphs[:20],
                "table_count": len(tables),
                "tables": tables
            }
        
        result = await loop.run_in_executor(None, extract_docx)
        
        summary = self._generate_summary(result["text"])
        keywords = self._extract_keywords(result["text"])
        
        return {
            "success": True,
            "text": result["text"][:50000],
            "summary": summary,
            "keywords": keywords,
            "paragraph_count": result["paragraph_count"],
            "paragraphs": result["paragraphs"],
            "table_count": result["table_count"],
            "tables": result["tables"]
        }
    
    async def _analyze_text(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Analyze plain text file"""
        text = content.decode("utf-8", errors="ignore")
        
        # Count lines and words
        lines = text.split("\n")
        words = text.split()
        
        summary = self._generate_summary(text)
        keywords = self._extract_keywords(text)
        
        return {
            "success": True,
            "text": text[:100000],  # Limit
            "summary": summary,
            "keywords": keywords,
            "line_count": len(lines),
            "word_count": len(words),
            "char_count": len(text)
        }
    
    async def _analyze_csv(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Analyze CSV file"""
        if not PANDAS_AVAILABLE:
            return {
                "success": False,
                "error": "Pandas not available for CSV analysis"
            }
        
        loop = asyncio.get_event_loop()
        
        def extract_csv():
            from io import BytesIO
            import pandas as pd
            
            csv_file = BytesIO(content)
            df = pd.read_csv(csv_file)
            
            return {
                "columns": list(df.columns),
                "row_count": len(df),
                "data": df.head(100).to_dict("records"),
                "statistics": df.describe().to_dict(),
                "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
        
        result = await loop.run_in_executor(None, extract_csv)
        
        return {
            "success": True,
            "text": f"CSV with {result['row_count']} rows and {result['columns']}",
            "summary": f"This CSV file contains {result['row_count']} rows and {len(result['columns'])} columns: {', '.join(result['columns'])}",
            "keywords": result["columns"][:10],
            "columns": result["columns"],
            "row_count": result["row_count"],
            "data": result["data"],
            "statistics": result["statistics"]
        }
    
    async def _analyze_excel(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Analyze Excel file"""
        if not PANDAS_AVAILABLE:
            return {
                "success": False,
                "error": "Pandas not available for Excel analysis"
            }
        
        loop = asyncio.get_event_loop()
        
        def extract_excel():
            from io import BytesIO
            import pandas as pd
            
            excel_file = BytesIO(content)
            
            # Get all sheet names
            xl = pd.ExcelFile(excel_file)
            sheets = {}
            
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheets[sheet_name] = {
                    "columns": list(df.columns),
                    "row_count": len(df),
                    "data": df.head(50).to_dict("records")
                }
            
            return sheets
        
        result = await loop.run_in_executor(None, extract_excel)
        
        summary = f"Excel file with {len(result)} sheets: {', '.join(result.keys())}"
        
        return {
            "success": True,
            "text": summary,
            "summary": summary,
            "keywords": list(result.keys()),
            "sheet_count": len(result),
            "sheets": result
        }
    
    async def _analyze_pptx(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Analyze PowerPoint file"""
        # Basic support - extract text
        return await self._analyze_text(content, filename)
    
    def _generate_summary(self, text: str, max_length: int = 500) -> str:
        """Generate a simple summary of text"""
        # Simple extractive summarization
        sentences = re.split(r'[.!?]+', text)
        
        if len(sentences) <= 3:
            return text[:max_length]
        
        # Take first few sentences as summary
        summary = ". ".join(sentences[:3]) + "."
        
        if len(summary) > max_length:
            summary = summary[:max_length] + "..."
        
        return summary
    
    def _extract_keywords(self, text: str, num_keywords: int = 20) -> List[str]:
        """Extract keywords from text"""
        # Simple keyword extraction
        # Remove common words
        stopwords = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to',
            'for', 'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are',
            'were', 'been', 'be', 'have', 'has', 'had', 'do', 'does', 'did',
            'will', 'would', 'could', 'should', 'may', 'might', 'must',
            'shall', 'can', 'need', 'this', 'that', 'these', 'those', 'i',
            'you', 'he', 'she', 'it', 'we', 'they', 'what', 'which', 'who',
            'when', 'where', 'why', 'how', 'all', 'each', 'every', 'both',
            'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor',
            'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 'just'
        }
        
        # Extract words
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        # Count frequencies
        word_freq = {}
        for word in words:
            if word not in stopwords:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Sort by frequency
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        
        # Return top keywords
        return [word for word, _ in sorted_words[:num_keywords]]
    
    async def answer_question(
        self,
        file_path: str,
        question: str,
        file_content: bytes = None,
        ai_client = None
    ) -> Dict[str, Any]:
        """
        Answer a question about a document
        
        Args:
            file_path: Path to file
            question: Question to answer
            file_content: Optional file content
            ai_client: AI client for generating answer
            
        Returns:
            Dict with answer and sources
        """
        # First analyze the file
        analysis = await self.analyze_file(file_path, file_content)
        
        if not analysis.get("success"):
            return analysis
        
        # Generate answer using AI
        if ai_client:
            context = analysis.get("text", "")[:8000]  # Limit context
            
            prompt = f"""Based on the following document, answer the question.

Document Content:
{context}

Question: {question}

Answer:"""
            
            try:
                # Use Groq for response
                from ai_aggregator import get_ai_aggregator
                aggregator = get_ai_aggregator()
                
                response = await aggregator.generate_response(
                    prompt=prompt,
                    provider="groq",
                    model="llama-3.1-70b-versatile"
                )
                
                return {
                    "success": True,
                    "answer": response,
                    "question": question,
                    "document": analysis.get("filename"),
                    "summary": analysis.get("summary")
                }
                
            except Exception as e:
                return {
                    "success": False,
                    "error": f"AI processing error: {str(e)}"
                }
        else:
            # Return extracted content for manual review
            return {
                "success": True,
                "answer": "Please provide AI configuration to get automatic answers.",
                "question": question,
                "document": analysis.get("filename"),
                "extracted_text": analysis.get("text", "")[:2000],
                "summary": analysis.get("summary")
            }


# Singleton instance
_document_analyzer: Optional[DocumentAnalyzer] = None


def get_document_analyzer() -> DocumentAnalyzer:
    """Get or create document analyzer singleton"""
    global _document_analyzer
    if _document_analyzer is None:
        _document_analyzer = DocumentAnalyzer()
    return _document_analyzer


async def analyze_document(
    file_path: str,
    file_content: bytes = None
) -> Dict[str, Any]:
    """
    Convenience function for document analysis
    """
    analyzer = get_document_analyzer()
    return await analyzer.analyze_file(file_path, file_content)


async def answer_document_question(
    file_path: str,
    question: str,
    file_content: bytes = None
) -> Dict[str, Any]:
    """
    Convenience function for document Q&A
    """
    analyzer = get_document_analyzer()
    return await analyzer.answer_question(file_path, question, file_content)
